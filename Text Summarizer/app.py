import streamlit as st
from io import BytesIO
from fpdf import FPDF
from docx import Document
from summarizer import summarize_text, extract_text_from_pdf, extract_text_from_docx, fetch_text_from_url

# Set page config
st.set_page_config(
    page_title="AI Text Summarizer",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom Styles for Enhanced UI
st.markdown("""
    <style>
        /* Main styling */
        .main {
            background-color: #f8f9fa;
        }
        
        /* Sidebar styling */
        section[data-testid="stSidebar"] {
            background-color: #2E4053;
            color: white;
        }
        
        section[data-testid="stSidebar"] .stRadio label {
            color: white !important;
        }
        
        /* Container styling */
        .block-container {
            padding: 2rem;
            max-width: 1200px;
        }
        
        /* Summary box styling */
        .summary-box {
            border: 1px solid #ddd; 
            padding: 20px; 
            border-radius: 10px; 
            max-height: 400px; 
            overflow-y: auto; 
            background-color: white; 
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            line-height: 1.6;
        }
        
        /* Header styling */
        h1, h2, h3 {
            color: #2E86C1;
        }
        
        /* Button styling */
        .stButton button {
            background-color: #2E4053;
            color: white;
            border-radius: 6px;
            padding: 10px 15px;
            font-weight: 600;
            transition: all 0.3s ease;
        }
        
        .stButton button:hover {
            background-color: #3E5063;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
        }
        
        /* Card styling for sections */
        .card {
            border-radius: 10px;
            padding: 20px;
            background-color: white;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
            margin-bottom: 20px;
        }
        
        /* Progress bar styling */
        .stProgress > div > div {
            background-color: #2E86C1;
        }
        
        /* Text area styling */
        textarea {
            border-radius: 8px !important;
            border: 1px solid #ddd !important;
        }
        
        /* File uploader styling */
        .stFileUploader > div {
            border-radius: 8px !important;
            padding: 10px !important;
        }
        
        /* Download button container */
        .download-container {
            display: flex;
            justify-content: space-around;
            margin-top: 15px;
        }
        
        /* Stats counter */
        .stats-counter {
            text-align: right;
            color: #666;
            font-size: 0.8em;
            margin-top: 5px;
        }
        
        /* Success message */
        .success-message {
            padding: 10px;
            background-color: #d4edda;
            color: #155724;
            border-radius: 5px;
            margin-bottom: 10px;
        }
    </style>
""", unsafe_allow_html=True)

# Create a sidebar for better organization
with st.sidebar:
    st.markdown("<h2 style='text-align: center; color: white;'>✍️ Input Options</h2>", unsafe_allow_html=True)
    input_type = st.radio("", ["Enter Text", "Upload PDF", "Upload DOCX", "Summarize URL"])
    
    st.markdown("<hr style='margin: 15px 0;'>", unsafe_allow_html=True)
    
    st.markdown("<h2 style='text-align: center; color: white;'>🎯 Summary Settings</h2>", unsafe_allow_html=True)
    summary_type = st.selectbox("📌 Summary Type", ["Concise", "Detailed", "Bullet Points"])
    length = st.select_slider("📏 Length", options=["Short", "Medium", "Long"])
    tone = st.selectbox("🎭 Tone", ["Neutral", "Formal", "Casual", "Storytelling", "Professional"])
    language = st.selectbox("🌐 Language", ["English", "Spanish", "French", "German"])
    
    st.markdown("<hr style='margin: 15px 0;'>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #AAAAAA; font-size: 0.8em;'>Powered by Advanced NLP</p>", unsafe_allow_html=True)

# App Title with Logo
st.markdown("<h1 style='text-align: center;'>📜 AI-Powered Text Summarizer</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; font-size: 1.2em;'>Transform long texts into concise, meaningful summaries</p>", unsafe_allow_html=True)

# Main content area with card-like container
st.markdown("<div class='card'>", unsafe_allow_html=True)

# Input handling
text = ""

if input_type == "Enter Text":
    text = st.text_area("📝 Enter the text to summarize:", height=200, 
                        placeholder="Paste your text here (minimum 100 characters for best results)...")

elif input_type == "Upload PDF":
    uploaded_file = st.file_uploader("📂 Upload a PDF file", type=["pdf"])
    if uploaded_file:
        with st.spinner("🔍 Extracting text from PDF..."):
            text = extract_text_from_pdf(uploaded_file)
        if text.startswith("Error:"):
            st.error(text)
        else:
            st.markdown("<div class='success-message'>✅ PDF successfully processed!</div>", unsafe_allow_html=True)
            with st.expander("View extracted text"):
                st.text_area("", text, height=200)

elif input_type == "Upload DOCX":
    uploaded_file = st.file_uploader("📂 Upload a DOCX file", type=["docx"])
    if uploaded_file:
        with st.spinner("🔍 Extracting text from DOCX..."):
            text = extract_text_from_docx(uploaded_file)
        if text.startswith("Error:"):
            st.error(text)
        else:
            st.markdown("<div class='success-message'>✅ DOCX successfully processed!</div>", unsafe_allow_html=True)
            with st.expander("View extracted text"):
                st.text_area("", text, height=200)

elif input_type == "Summarize URL":
    url = st.text_input("🌍 Enter a URL:", placeholder="https://example.com/article")
    col1, col2 = st.columns([3, 1])
    with col2:
        fetch_button = st.button("🔎 Fetch Content", use_container_width=True)
    
    if fetch_button and url:
        with st.spinner("🔍 Fetching and analyzing web content..."):
            text = fetch_text_from_url(url)
        if text.startswith("Error:"):
            st.error(text)
        else:
            st.markdown("<div class='success-message'>✅ Content successfully fetched!</div>", unsafe_allow_html=True)
            text_length = len(text.split())
            st.info(f"📊 Retrieved {text_length} words from the URL")
            with st.expander("View extracted text"):
                st.text_area("", text, height=200)

# Display word count if text exists
if text.strip():
    st.markdown(f"<p class='stats-counter'>{len(text.split())} words | {len(text)} characters</p>", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)

# Summarize Button - only show if we have text
if text.strip() and len(text.split()) > 10:
    summarize_button = st.button("⚡ Generate Summary", use_container_width=True)
    
    if summarize_button:
        summary = ""
        with st.spinner("⏳ Generating summary..."):
            # Add progress bar for better UX
            progress_bar = st.progress(0)
            for i in range(100):
                # Simulating progress for UX purposes
                import time
                time.sleep(0.01)
                progress_bar.progress(i + 1)
            
            # Actually perform the summarization
            summary = summarize_text(text, summary_type, length, tone, language)
        
        # Remove progress bar after completion
        progress_bar.empty()
        
        if not summary.startswith("Error:"):
            # Display summary inside a scrollable block
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("### 📌 Summary:")
            
            # Fix for the f-string backslash issue
            formatted_summary = summary.replace("\n", "<br>")
            st.markdown(
                f'<div class="summary-box">{formatted_summary}</div>',
                unsafe_allow_html=True
            )
            
            # Display summary stats
            original_length = len(text.split())
            summary_length = len(summary.split())
            reduction = round((1 - summary_length / original_length) * 100)
            
            st.markdown(f"""
                <p class='stats-counter'>
                    Original: {original_length} words | Summary: {summary_length} words | {reduction}% reduction
                </p>
            """, unsafe_allow_html=True)
            
            # Download options
            st.markdown("### 📥 Download Summary")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.download_button("📄 TXT", summary, file_name="summary.txt", use_container_width=True)
            
            with col2:
                docx_data = BytesIO()
                doc = Document()
                doc.add_heading("Text Summary", 0)
                doc.add_paragraph(summary)
                doc.add_paragraph(f"\nSummary generated using AI Text Summarizer")
                doc.save(docx_data)
                docx_data.seek(0)
                st.download_button("📜 DOCX", data=docx_data, file_name="summary.docx",
                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                 use_container_width=True)
            
            with col3:
                pdf_data = BytesIO()
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                
                # Fix the encoding issue by handling utf-8 characters
                try:
                    pdf.multi_cell(190, 10, summary)
                    pdf_bytes = pdf.output(dest='S').encode('latin-1')
                except UnicodeEncodeError:
                    # Fallback to basic ASCII if encoding fails
                    pdf.multi_cell(190, 10, summary.encode('ascii', 'replace').decode('ascii'))
                    pdf_bytes = pdf.output(dest='S').encode('latin-1')
                    
                pdf_data.write(pdf_bytes)
                pdf_data.seek(0)
                
                st.download_button("📕 PDF", data=pdf_data, file_name="summary.pdf", 
                                  mime="application/pdf", use_container_width=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.error(summary)
    
else:
    if text.strip() and len(text.split()) <= 10:
        st.warning("⚠ Please enter more text for a meaningful summary (at least 10 words).")
    elif "url" in locals() and url and not text:
        st.info("👆 Enter a URL and click 'Fetch Content' to get started.")
    elif "uploaded_file" in locals() and not uploaded_file and not text:
        st.info("👆 Upload a document to get started.")
    elif not text:
        st.info("👆 Enter some text, upload a document, or provide a URL to get started.")