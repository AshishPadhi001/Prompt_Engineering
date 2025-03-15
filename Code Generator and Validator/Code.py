from transformers import pipeline
from docx import Document
from fpdf import FPDF

# **Load Hugging Face LLM optimized for code generation**
code_generator = pipeline("text-generation", model="Salesforce/codegen-350M-mono")

# **Supported Programming Languages & File Extensions**
LANGUAGE_EXTENSIONS = {
    "python": ".py",
    "java": ".java",
    "c++": ".cpp",
    "c": ".c",
    "c#": ".cs",
    "javascript": ".js",
    "typescript": ".ts",
    "ruby": ".rb",
    "php": ".php",
    "swift": ".swift",
    "go": ".go",
    "rust": ".rs",
    "r": ".r",
    "julia": ".jl",
    "dart": ".dart",
    "objective-c": ".m",
    "objective-c++": ".mm",
    "scala": ".scala",
    "f#": ".fs",
    "haskell": ".hs",
    "elixir": ".ex",
    "erlang": ".erl",
    "clojure": ".clj",
    "scheme": ".scm",
    "lisp": ".lisp",
    "prolog": ".pl",
    "tcl": ".tcl",
    "bash": ".sh",
    "shell": ".sh",
    "powershell": ".ps1",
    "html": ".html",
    "css": ".css",
    "scss": ".scss",
    "sass": ".sass",
    "sql": ".sql",
    "plsql": ".pls",
    "graphql": ".graphql",
    "xml": ".xml",
    "json": ".json",
    "yaml": ".yaml",
    "toml": ".toml",
    "markdown": ".md",
    "latex": ".tex",
    "rest": ".rst",
    "matlab": ".m",
    "jupyter": ".ipynb",
    "assembly": ".asm",
    "vhdl": ".vhdl",
    "verilog": ".v",
    "cobol": ".cbl",
    "fortran": ".f90",
    "ada": ".adb",
    "groovy": ".groovy",
    "kotlin": ".kt",
    "lua": ".lua",
    "perl": ".pl",
    "abap": ".abap",
    "actionscript": ".as",
    "apex": ".cls",
    "awk": ".awk",
    "bcpl": ".bcpl",
    "blitzmax": ".bmx",
    "boo": ".boo",
    "ceylon": ".ceylon",
    "chapel": ".chpl",
    "clean": ".icl",
    "crystal": ".cr",
    "delphi": ".pas",
    "eiffel": ".e",
    "elm": ".elm",
    "factor": ".factor",
    "forth": ".fs",
    "gams": ".gms",
    "gap": ".g",
    "idris": ".idr",
    "io": ".io",
    "j": ".ijs",
    "janet": ".janet",
    "modula-2": ".mod",
    "nim": ".nim",
    "ocaml": ".ml",
    "opencl": ".cl",
    "pike": ".pike",
    "postscript": ".ps",
    "rebol": ".r",
    "rex": ".rex",
    "ring": ".ring",
    "sml": ".sml",
    "smalltalk": ".st",
    "spice": ".cir",
    "stan": ".stan",
    "turing": ".t",
    "vala": ".vala",
    "visual basic": ".vb",
    "x10": ".x10",
    "xtend": ".xtend",
    "zig": ".zig",
    "terraform": ".tf",
    "dockerfile": "Dockerfile",
    "kubernetes": ".yaml",
    "ansible": ".yml",
    "hcl": ".hcl",
    "nix": ".nix",
    "racket": ".rkt",
    "red": ".red",
    "pony": ".pony"
}


def generate_code(question, language):
    """
    Generates optimized, well-commented code for a given problem.
    """
    prompt = (
        f"Write a {language} program to solve the following problem:\n\n"
        f"### Problem Description:\n{question}\n\n"
        f"### Requirements:\n"
        f"- Follow best coding practices\n"
        f"- Include meaningful comments\n"
        f"- Ensure readability and efficiency\n\n"
        f"### Code:\n"
    )

    try:
        response = code_generator(prompt, max_length=700, temperature=0.5, top_p=0.9, do_sample=True)
        generated_code = response[0]["generated_text"].strip()
        return generated_code
    except Exception as e:
        return f"❌ Error: {str(e)}"

def explain_code(code, language):
    """
    Generates a step-by-step explanation of the provided code.
    """
    prompt = (
        f"Explain the following {language} code in a structured, step-by-step manner:\n\n"
        f"### Code:\n{code}\n\n"
        f"### Explanation:\n"
    )

    try:
        response = code_generator(prompt, max_length=700, temperature=0.5, top_p=0.9, do_sample=True)
        explanation = response[0]["generated_text"].strip()
        return explanation
    except Exception as e:
        return f"❌ Error: {str(e)}"

def save_file(text, language):
    """
    Saves generated code or explanations with the correct file format.
    """
    file_name = input("\n📁 Enter file name (without extension): ").strip()
    ext = LANGUAGE_EXTENSIONS.get(language.lower(), ".txt")

    print("\n💾 Choose file format:")
    print(f"1️⃣  Save as {ext} (Source Code File)")
    print("2️⃣  Save as DOCX")
    print("3️⃣  Save as PDF")
    print("4️⃣  Skip saving")

    choice = input("\nEnter your choice: ").strip()

    if choice == "1":
        with open(f"{file_name}{ext}", "w", encoding="utf-8") as file:
            file.write(text)
        print(f"\n✅ Code saved as '{file_name}{ext}'!")

    elif choice == "2":
        doc = Document()
        doc.add_heading("Generated Code", 0)
        doc.add_paragraph(text)
        doc.save(f"{file_name}.docx")
        print(f"\n✅ Code saved as '{file_name}.docx'!")

    elif choice == "3":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(190, 10, text)
        pdf.output(f"{file_name}.pdf")
        print(f"\n✅ Code saved as '{file_name}.pdf'!")

    elif choice == "4":
        print("\n🚀 Code was not saved. Returning to menu.")
    else:
        print("❌ Invalid choice. Please try again.")

def main():
    """
    Main function to interact with the AI Code Generator & Explainer.
    """
    print("\n🤖 Welcome to the AI Code Generator & Explainer!")

    while True:
        print("\n📌 Choose an option:")
        print("1️⃣  Generate Code")
        print("2️⃣  Explain Code")
        print("3️⃣  Exit")

        choice = input("\nEnter your choice: ").strip()

        if choice == "1":
            question = input("\n📝 Enter your coding problem: ").strip()
            language = input("💻 Enter the programming language: ").strip()

            print("\n⏳ Generating your code...")
            generated_code = generate_code(question, language)

            print("\n✨ **Generated Code:**\n")
            print(generated_code)

            save_choice = input("\n💾 Do you want to save this code? (yes/no): ").strip().lower()
            if save_choice == "yes":
                save_file(generated_code, language)

        elif choice == "2":
            language = input("\n💻 Enter the programming language of the code: ").strip()
            print("\n✍️ Paste your code below and press Enter when done:")
            user_code = []
            while True:
                line = input()
                if line.strip() == "":
                    break
                user_code.append(line)
            user_code = "\n".join(user_code)

            print("\n⏳ Generating explanation...")
            explanation = explain_code(user_code, language)

            print("\n📝 **Code Explanation:**\n")
            print(explanation)

            save_choice = input("\n💾 Do you want to save this explanation? (yes/no): ").strip().lower()
            if save_choice == "yes":
                save_file(explanation, language)

        elif choice == "3":
            print("\n🚀 Exiting the program. Have a great day!\n")
            break

        else:
            print("❌ Invalid choice. Please enter a valid option.")

if __name__ == "__main__":
    main()
