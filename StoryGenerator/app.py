import streamlit as st
import os
import time
from PIL import Image
import io
import base64

# Import functions from story.py
try:
    from story import generate_story
except ImportError:
    st.error("Could not import from story.py. Make sure the file is in the same directory.")
    st.stop()

def save_story_to_file(text, file_format, filename):
    """
    Save the story to a file with the specified format
    """
    if not filename:
        return False, "Please provide a filename"
    
    try:
        if file_format == "txt":
            with open(f"{filename}.txt", "w", encoding="utf-8") as file:
                file.write(text)
            return True, f"Story saved as '{filename}.txt'"
            
        elif file_format == "md":
            with open(f"{filename}.md", "w", encoding="utf-8") as file:
                file.write(f"# {filename}\n\n{text}")
            return True, f"Story saved as '{filename}.md'"
            
        # For DOCX and PDF formats, we'll import the libraries here to avoid unnecessary imports
        elif file_format == "docx":
            from docx import Document
            doc = Document()
            doc.add_heading(filename, 0)
            doc.add_paragraph(text)
            doc.save(f"{filename}.docx")
            return True, f"Story saved as '{filename}.docx'"
            
        elif file_format == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(190, 10, text)
            pdf.output(f"{filename}.pdf")
            return True, f"Story saved as '{filename}.pdf'"
    
    except Exception as e:
        return False, f"Error saving file: {str(e)}"

def get_download_link(text, filename):
    """Generate a link to download the text as a file"""
    b64 = base64.b64encode(text.encode()).decode()
    return f'<a href="data:text/plain;base64,{b64}" download="{filename}.txt">Download story as TXT</a>'

def main():
    st.set_page_config(
        page_title="TaleCraft AI",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS styling
    st.markdown("""
        <style>
        .main-header {
            font-size: 2.5rem;
            color: #4a6ea9;
            margin-bottom: 0.5rem;
        }
        .sub-header {
            font-size: 1.2rem;
            color: #666;
            margin-bottom: 2rem;
        }
        .story-box {
            background-color: #f8f9fa;
            border-radius: 10px;
            padding: 20px;
            border: 1px solid #ddd;
        }
        .story-title {
            font-size: 1.5rem;
            color: #333;
            margin-bottom: 1rem;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown('<h1 class="main-header">📜 AI Story Generator</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Create captivating stories with the help of AI</p>', unsafe_allow_html=True)
    
    # Sidebar for mode selection
    st.sidebar.title("Mode Selection")
    app_mode = st.sidebar.radio("Choose Mode", ["Generate Full Story", "Interactive Story Mode"])
    
    # About section in sidebar
    with st.sidebar.expander("About this app"):
        st.write("""
        This app uses AI to generate stories based on your inputs. 
        Choose from two modes:
        - **Generate Full Story**: AI creates a complete story
        - **Interactive Story Mode**: AI starts, you continue
        
        Made with Streamlit and Hugging Face's transformers library.
        """)
    
    if app_mode == "Generate Full Story":
        st.header("Generate a Complete Story")
        
        # Input form
        with st.form("story_form"):
            col1, col2 = st.columns(2)
            
            with col1:
                theme = st.text_input("Theme (e.g., Adventure, Friendship, Betrayal)", "Adventure")
                genre = st.text_input("Genre (e.g., Fantasy, Sci-Fi, Horror, Romance)", "Fantasy")
            
            with col2:
                input_words = st.text_input("Keywords or Phrase to Guide the Story", "magic forest hidden treasure")
                length = st.slider("Story Length", min_value=50, max_value=1000, value=300, step=50, 
                                 help="Minimum 50 words, recommended 150-500")
            
            generate_button = st.form_submit_button("Generate Story")
            
        if generate_button:
            with st.spinner("✨ Creating your story..."):
                # Call the generate_story function from story.py
                story = generate_story(theme, genre, input_words, length)
                
                # Store the story in session state to keep it when regenerating
                st.session_state.current_story = story
        
        # Display story if available
        if 'current_story' in st.session_state:
            st.markdown('<div class="story-box">', unsafe_allow_html=True)
            st.markdown(f'<div class="story-title">A Journey of {st.session_state.get("theme", theme)}</div>', unsafe_allow_html=True)
            st.write(st.session_state.current_story)
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Actions after story generation
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("Regenerate Story"):
                    with st.spinner("🔄 Regenerating your story..."):
                        story = generate_story(theme, genre, input_words, length)
                        st.session_state.current_story = story
                        st.experimental_rerun()
            
            # Save options
            with col2:
                save_format = st.selectbox("Save Format", ["txt", "docx", "md", "pdf"])
            
            filename = st.text_input("Filename (without extension)")
            
            if st.button("Save Story"):
                if filename:
                    success, message = save_story_to_file(st.session_state.current_story, save_format, filename)
                    if success:
                        st.success(message)
                    else:
                        st.error(message)
                else:
                    st.warning("Please provide a filename")
            
            # Download directly in browser
            st.markdown(get_download_link(st.session_state.current_story, filename or "story"), unsafe_allow_html=True)
    
    else:  # Interactive Story Mode
        st.header("Interactive Story Mode")
        st.write("AI will write part of the story, and you will continue it!")
        
        # Input form
        with st.form("interactive_form"):
            col1, col2 = st.columns(2)
            
            with col1:
                theme = st.text_input("Theme", "Mystery")
                genre = st.text_input("Genre", "Thriller")
            
            with col2:
                input_words = st.text_input("Keywords or Phrase to Start", "abandoned mansion dark night")
            
            start_button = st.form_submit_button("Start Story")
        
        if start_button:
            with st.spinner("✨ Creating story opening..."):
                # Generate the initial part of the story
                initial_story = generate_story(theme, genre, input_words, 150)
                st.session_state.initial_story = initial_story
                st.session_state.theme = theme
        
        # Display initial story and allow user to continue
        if 'initial_story' in st.session_state:
            st.markdown('<div class="story-box">', unsafe_allow_html=True)
            st.markdown(f'<div class="story-title">A Journey of {st.session_state.theme}</div>', unsafe_allow_html=True)
            st.write(st.session_state.initial_story)
            st.markdown('</div>', unsafe_allow_html=True)
            
            st.write("✍️ Now it's your turn! Continue the story:")
            user_contribution = st.text_area("Your contribution", height=150)
            
            if st.button("Complete Story"):
                if user_contribution:
                    full_story = st.session_state.initial_story + "\n\n" + user_contribution
                    st.session_state.full_story = full_story
                else:
                    st.warning("Please add your contribution to the story")
        
        # Display and save options for the complete story
        if 'full_story' in st.session_state:
            st.subheader("Complete Story")
            st.markdown('<div class="story-box">', unsafe_allow_html=True)
            st.write(st.session_state.full_story)
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Save options
            save_format = st.selectbox("Save Format", ["txt", "docx", "md", "pdf"])
            filename = st.text_input("Filename (without extension)")
            
            if st.button("Save Complete Story"):
                if filename:
                    success, message = save_story_to_file(st.session_state.full_story, save_format, filename)
                    if success:
                        st.success(message)
                    else:
                        st.error(message)
                else:
                    st.warning("Please provide a filename")
            
            # Download directly in browser
            st.markdown(get_download_link(st.session_state.full_story, filename or "interactive_story"), unsafe_allow_html=True)

if __name__ == "__main__":
    main()