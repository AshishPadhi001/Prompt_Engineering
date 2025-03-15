from transformers import pipeline
from docx import Document
from fpdf import FPDF

# Load a Hugging Face model for story generation
story_generator = pipeline("text-generation", model="distilgpt2")

def generate_story(theme: str, genre: str, input_words: str, length: int) -> str:
    """
    Generates a structured short story based on the given theme, genre, and input words.

    :param theme: The central theme of the story (e.g., "Adventure", "Friendship", "Betrayal").
    :param genre: The story genre (e.g., "Fantasy", "Sci-Fi", "Horror").
    :param input_words: Key words or a sentence to guide the story.
    :param length: Maximum length of the generated story (user-defined).
    :return: The generated short story.
    """
    if not theme or not genre or not input_words or length < 50:
        return "Error: Please provide a valid theme, genre, input words, and a minimum length of 50 words."

    # Optimized storytelling prompt
    prompt = (
        f"Write a captivating {genre.lower()} story with the theme of {theme.lower()}. "
        f"Ensure the story has a structured format: an engaging opening, a gripping middle, and a satisfying conclusion. "
        f"Incorporate the following elements: {input_words}. "
        f"Use rich descriptions, deep character development, and immersive world-building.\n\n"
        f"Title: A Journey of {theme}\n\n"
        f"Once upon a time, "
    )

    try:
        story = story_generator(prompt, max_length=length, temperature=0.8, top_p=0.95, do_sample=True)
        return story[0]["generated_text"]
    except Exception as e:
        return f"Error generating story: {str(e)}"

def save_story(text: str):
    """
    Allows the user to choose how they want to save the story.

    :param text: The generated story.
    """
    while True:
        print("\n📥 **Save Story Options**")
        print("1️⃣  Save as TXT")
        print("2️⃣  Save as DOCX")
        print("3️⃣  Save as Markdown (MD)")
        print("4️⃣  Save as PDF")
        print("5️⃣  Skip saving")

        choice = input("\nEnter your choice: ").strip()

        if choice in ["1", "2", "3", "4"]:
            file_name = input("\nEnter file name (without extension): ").strip()
            
            if choice == "1":
                with open(f"{file_name}.txt", "w", encoding="utf-8") as file:
                    file.write(text)
                print(f"\n✅ Story saved as '{file_name}.txt'!")
                
            elif choice == "2":
                doc = Document()
                doc.add_heading(file_name, 0)
                doc.add_paragraph(text)
                doc.save(f"{file_name}.docx")
                print(f"\n✅ Story saved as '{file_name}.docx'!")
                
            elif choice == "3":
                with open(f"{file_name}.md", "w", encoding="utf-8") as file:
                    file.write(f"# {file_name}\n\n{text}")
                print(f"\n✅ Story saved as '{file_name}.md'!")
                
            elif choice == "4":
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(190, 10, text)
                pdf.output(f"{file_name}.pdf")
                print(f"\n✅ Story saved as '{file_name}.pdf'!")
            
            break  # Exit loop after saving

        elif choice == "5":
            print("\n🚀 Story was not saved. Returning to menu.")
            break  # Skip saving and return to menu
        else:
            print("❌ Invalid choice. Please select a valid option.")

def interactive_story_mode():
    """
    Interactive mode where AI writes part of the story, and the user contributes to it.
    """
    print("\n📝 **Interactive Story Mode**: AI will write part of the story, and you will continue it!")
    theme = input("\nEnter a theme: ").strip()
    genre = input("Enter a genre: ").strip()
    input_words = input("Enter key words or a phrase to start: ").strip()
    
    print("\n⏳ Generating story opening...\n")
    initial_story = generate_story(theme, genre, input_words, 150)
    print("\n✨ **AI's Story Start:**\n")
    print(initial_story)
    
    user_contribution = input("\n✍️ Now it's your turn! Continue the story: ").strip()
    full_story = initial_story + "\n\n" + user_contribution

    save_story(full_story)

def display_menu():
    """
    Displays the main menu and handles user input.
    """
    while True:
        print("\n📜 **AI Story Generator**")
        print("1️⃣  Generate a full story")
        print("2️⃣  Interactive story mode")
        print("3️⃣  Exit")

        choice = input("\nEnter your choice: ").strip()

        if choice == "1":
            theme = input("\nEnter a theme for the story (e.g., Adventure, Friendship, Betrayal): ").strip()
            genre = input("Enter the genre (e.g., Fantasy, Sci-Fi, Horror, Romance): ").strip()
            input_words = input("Enter key words or a sentence to guide the story: ").strip()
            
            while True:
                try:
                    length = int(input("Enter the story length (minimum 50 words, recommended 150-500): ").strip())
                    if length < 50:
                        print("⚠ Story length must be at least 50 words.")
                        continue
                    break
                except ValueError:
                    print("❌ Invalid input. Please enter a valid number.")

            while True:
                print("\n⏳ Generating your story...\n")
                story = generate_story(theme, genre, input_words, length)
                print("\n✨ **Generated Story:**\n")
                print(story)

                regenerate_choice = input("\n🔄 Do you want to regenerate the story? (yes/no): ").strip().lower()
                if regenerate_choice == "no":
                    save_story(story)
                    break
                elif regenerate_choice == "yes":
                    print("\n🔄 Regenerating story...\n")
                else:
                    print("❌ Invalid choice. Please enter 'yes' or 'no'.")

        elif choice == "2":
            interactive_story_mode()

        elif choice == "3":
            print("\n🚀 Exiting the program. Have a great day!\n")
            break

        else:
            print("❌ Invalid choice. Please enter a valid option.")

# Run the program
if __name__ == "__main__":
    display_menu()
