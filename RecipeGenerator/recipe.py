from transformers import pipeline
from docx import Document
from fpdf import FPDF

# **Load Hugging Face Model (Fast & Efficient)**
MODEL_NAME = "gpt2-medium"
recipe_generator = pipeline("text-generation", model=MODEL_NAME)

def generate_recipe(ingredients, diet, meal_type, cooking_time, servings):
    """
    Generates a detailed recipe based on user input using an optimized prompt.
    """
    prompt = (
        f"Write a detailed {meal_type} recipe using these ingredients: {ingredients}. "
        f"The recipe should follow a {diet} diet, take around {cooking_time} minutes, and serve {servings} people.\n\n"
        f"### Recipe Format:\n"
        f"- **Title**: A unique and appealing name.\n"
        f"- **Ingredients**: Clearly list all required ingredients.\n"
        f"- **Preparation Steps**: Provide detailed step-by-step cooking instructions.\n"
        f"- **Chef’s Tips**: Share cooking hacks or variations.\n"
        f"- **Nutritional Information**: Calories, protein, carbs, fats.\n"
    )

    try:
        response = recipe_generator(prompt, max_length=400, temperature=0.7, top_p=0.9)
        return response[0]["generated_text"].strip()
    except Exception as e:
        return f"❌ Error: {str(e)}"

def get_dish_details(dish_name):
    """
    Fetches details for a specific dish.
    """
    prompt = f"Write a complete recipe for {dish_name}."

    try:
        response = recipe_generator(prompt, max_length=400, temperature=0.7, top_p=0.9)
        return response[0]["generated_text"].strip()
    except Exception as e:
        return f"❌ Error: {str(e)}"

def save_recipe(text):
    """
    Saves the generated recipe in different formats.
    """
    file_name = input("\n📁 Enter recipe name (without extension): ").strip()

    print("\n💾 Choose file format:")
    print("1️⃣  Save as TXT")
    print("2️⃣  Save as DOCX")
    print("3️⃣  Save as PDF")
    print("4️⃣  Skip saving")

    choice = input("\nEnter your choice: ").strip()

    if choice == "1":
        with open(f"{file_name}.txt", "w", encoding="utf-8") as file:
            file.write(text)
        print(f"\n✅ Recipe saved as '{file_name}.txt'!")

    elif choice == "2":
        doc = Document()
        doc.add_heading("Generated Recipe", 0)
        doc.add_paragraph(text)
        doc.save(f"{file_name}.docx")
        print(f"\n✅ Recipe saved as '{file_name}.docx'!")

    elif choice == "3":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(190, 10, text)
        pdf.output(f"{file_name}.pdf")
        print(f"\n✅ Recipe saved as '{file_name}.pdf'!")

    elif choice == "4":
        print("\n🚀 Recipe was not saved. Returning to menu.")
    else:
        print("❌ Invalid choice. Please try again.")

def main():
    """
    Main function to interact with the AI Recipe Generator.
    """
    print(f"\n🍽️ Welcome to the AI Recipe Generator!\nUsing free Hugging Face model: {MODEL_NAME}")

    while True:
        print("\n📌 Choose an option:")
        print("1. Generate a Recipe from Ingredients")
        print("2. Get Details for a Specific Dish")
        print("3. Exit")

        choice = input("\nEnter your choice (1-3): ").strip()

        if choice == "1":
            ingredients = input("\n🥕 Enter the ingredients you have (comma-separated): ").strip()
            diet = input("🥗 Enter your dietary preference (e.g., Vegan, Keto, Gluten-Free, None): ").strip()
            meal_type = input("🍲 What type of meal? (Breakfast, Lunch, Dinner, Snack, Dessert): ").strip()
            cooking_time = input("⏳ How much time do you have? (in minutes): ").strip()
            servings = input("👨‍👩‍👧‍👦 How many servings do you need? ").strip()

            print("\n⏳ Generating your personalized recipe...")
            recipe = generate_recipe(ingredients, diet, meal_type, cooking_time, servings)

            print("\n✨ **Generated Recipe:**\n")
            print(recipe)

            save_choice = input("\n💾 Do you want to save this recipe? (yes/no): ").strip().lower()
            if save_choice == "yes":
                save_recipe(recipe)

        elif choice == "2":
            dish_name = input("\n🍽️ Enter the name of the dish you want to learn about: ").strip()
            print(f"\n⏳ Generating detailed recipe for {dish_name}...\n")

            dish_details = get_dish_details(dish_name)
            print("\n✨ **Dish Details:**\n")
            print(dish_details)

            save_choice = input("\n💾 Do you want to save this recipe? (yes/no): ").strip().lower()
            if save_choice == "yes":
                save_recipe(dish_details)

        elif choice == "3":
            print("\n🚀 Exiting the program. Enjoy your meal!\n")
            break

        else:
            print("❌ Invalid choice. Please enter a valid option.")

if __name__ == "__main__":
    main()
